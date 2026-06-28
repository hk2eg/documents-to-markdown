import argparse
import docx
import re
import os
import pypdfium2 as pdfium

def extract_text_from_pdf(pdf_path):
    text = ""
    pdf = pdfium.PdfDocument(pdf_path)
    for page in pdf:
        textpage = page.get_textpage()
        text += textpage.get_text_range() + "\n"
    return text

def verify(source_path, md_path, keywords):
    print(f"Starting integrity verification for {md_path} against {source_path}...")
    
    if not os.path.exists(source_path):
        print(f"Error: Source file {source_path} not found.")
        return
    if not os.path.exists(md_path):
        print(f"Error: Output file {md_path} not found.")
        return

    ext = os.path.splitext(source_path)[1].lower()
    
    # 1. Inspect original document
    orig_chars = 0
    orig_words = 0
    orig_tables_count = 0
    
    if ext == '.docx':
        doc = docx.Document(source_path)
        docx_text = "\n".join([p.text for p in doc.paragraphs])
        orig_tables_count = len(doc.tables)
        orig_chars = len(docx_text)
        orig_words = len(docx_text.split())
    elif ext == '.pdf':
        pdf_text = extract_text_from_pdf(source_path)
        orig_chars = len(pdf_text)
        orig_words = len(pdf_text.split())
        # Hard to count tables reliably in PDF without full parser, set to 0 to skip check
        orig_tables_count = 0
    else:
        print(f"Warning: Unsupported source format {ext} for detailed metrics.")

    # 2. Inspect generated Markdown
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()
    
    gen_chars = len(md_content)
    gen_words = len(md_content.split())
    
    md_table_count = len(re.findall(r'\n\|[-\s:|]+\|\n', md_content))
    
    print("\n--- Structural Comparison ---")
    print(f"{'Metric':<20} | {'Original':<10} | {'Converted':<10} | {'Status'}")
    print("-" * 60)
    
    char_status = "PASS" if gen_chars > (orig_chars * 0.5) else "WARN (Low char count)"
    print(f"{'Characters':<20} | {orig_chars:<10} | {gen_chars:<10} | {char_status}")
    
    if ext == '.docx':
        table_status = "PASS" if md_table_count >= orig_tables_count else f"WARN ({orig_tables_count - md_table_count} tables may be missing)"
        print(f"{'Tables':<20} | {orig_tables_count:<10} | {md_table_count:<10} | {table_status}")
    
    # 3. Keyword Content Check
    if keywords:
        print("\n--- Content Integrity Check ---")
        missing_keywords = []
        for kw in keywords:
            if kw.lower() in md_content.lower():
                print(f"- Keyword '{kw}': FOUND")
            else:
                print(f"- Keyword '{kw}': NOT FOUND")
                missing_keywords.append(kw)
                
        if not missing_keywords:
            print("\nAll core keywords found. Content appears intact.")
        else:
            print(f"\nWarning: Missing keywords: {', '.join(missing_keywords)}")

    print("\nVerification complete.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Verify integrity of converted Markdown.")
    parser.add_argument("source_doc", help="Original source document (PDF, DOCX)")
    parser.add_argument("output_md", help="Converted Markdown file")
    parser.add_argument("--keywords", nargs="+", default=[], help="Keywords to verify in the output")
    
    args = parser.parse_args()
    verify(args.source_doc, args.output_md, args.keywords)
